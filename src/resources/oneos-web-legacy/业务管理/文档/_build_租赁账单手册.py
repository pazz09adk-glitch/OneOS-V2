# -*- coding: utf-8 -*-
"""生成「租赁账单」用户说明书 Word，输出到用户桌面。"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_cn(doc, text, level):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "PingFang SC"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    return h


def add_para_cn(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "PingFang SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    run.bold = bold
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = "PingFang SC"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def build(out_path: Path):
    doc = Document()
    sect = doc.sections[0]
    sect.left_margin = Cm(2.2)
    sect.right_margin = Cm(2.2)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("数字化资产 ONEOS 运管平台\n租赁账单 · 功能说明（用户操作）")
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "PingFang SC"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    add_para_cn(
        doc,
        "本文档依据原型页面「租赁账单.jsx」「租赁账单-查看.jsx」「租赁账单-收费明细.jsx」整理，"
        "说明从列表进入查看、再进入收费明细的常规操作与字段含义。正式系统以线上版本为准。",
    )

    add_heading_cn(doc, "一、模块概述", 1)
    add_para_cn(
        doc,
        "「租赁账单」用于按租赁合同聚合展示多期账单，支持筛选、展开查看每期账单明细与金额，"
        "在账单未提交前维护收费明细及列表中的部分成本字段；提交后收费明细不可再改。",
    )

    add_heading_cn(doc, "二、如何进入", 1)
    add_bullet(doc, "登录系统后，左侧菜单选择「业务管理」→「租赁账单」。")
    add_bullet(doc, "列表页面包屑为：业务管理 / 租赁账单。右上角可点「查看需求说明」打开产品需求全文。")

    add_heading_cn(doc, "三、列表页操作", 1)

    add_heading_cn(doc, "3.1 筛选区", 2)
    add_para_cn(doc, "第一行默认展示：合同编码、项目名称、客户名称（均为可搜索下拉，支持输入关键字匹配）。")
    add_para_cn(doc, "点击「展开」后增加：业务部门、业务负责人（下拉选择）、交车任务编码（文本框，模糊匹配）。")
    add_bullet(doc, "重置：清空所有筛选条件。")
    add_bullet(doc, "查询：按条件刷新列表（原型中与筛选联动）。")

    add_heading_cn(doc, "3.2 主表（合同维度）", 2)
    add_para_cn(
        doc,
        "每行代表一份租赁合同。点击行首「+」展开子表，查看该合同下各期账单。字段包括："
        "合同编码、合同类型（正式/试用）、项目名称、客户名称、合同生效日期、交车任务编码、业务部门、业务负责人。",
    )
    add_para_cn(doc, "表格底部分页：可切换页码、每页条数，并显示总条数。")

    add_heading_cn(doc, "3.3 子表（账单期维度）", 2)
    add_para_cn(doc, "子表列说明（与原型一致）：")
    tbl = doc.add_table(rows=13, cols=2)
    tbl.style = "Table Grid"
    hdr = ["字段", "说明"]
    for j, h in enumerate(hdr):
        tbl.rows[0].cells[j].text = h
        set_cell_shading(tbl.rows[0].cells[j], "E6F4EA")
    rows = [
        ("账单编号", "合同编码 + ZD + 四位期号。例：HT-ZL-2025-001 第1期 → …ZD0001。用于后期与用友 YS 等对账取票。"),
        ("账单期数", "该合同下第几期账单。"),
        ("状态", "已提交 / 待提交等。需求中还规划「已结清」（财务到账等于实收后）。"),
        ("账单开始/结束日期", "本期计费区间，格式 YYYY-MM-DD。"),
        ("提车数量", "显示为「N辆」链接，点击弹出车辆列表：品牌、型号、车牌号。"),
        ("应收款总额", "系统计算的应收金额（需求：各车月租金总和 + 各车服务费总和）。"),
        ("实收款总额", "需求：月租金+服务费合计减去减免总金额。"),
        ("减免总金额", "各减免项合计。"),
        ("车辆成本", "按车型日成本 × 账单天数 × 车辆数汇总展示（只读计算）。"),
        ("氢费成本", "点击金额可改为输入框，两位小数，失焦保存（列表内快速维护）。"),
        ("其他成本", "同上，可点击编辑。"),
        ("操作", "见下文。"),
    ]
    for i, (a, b) in enumerate(rows, start=1):
        tbl.rows[i].cells[0].text = a
        tbl.rows[i].cells[1].text = b

    add_heading_cn(doc, "3.4 列表行操作：查看 / 收费明细", 2)
    add_bullet(doc, "查看：任意状态均可进入「租赁账单 - 查看」页，浏览该期账单的只读明细与汇总。")
    add_bullet(
        doc,
        "收费明细：仅当该期状态为「待提交」时显示；「已提交」后入口隐藏，表示收费项已确认不可再改。",
    )

    add_heading_cn(doc, "四、查看页（只读）", 1)
    add_para_cn(doc, "面包屑：业务管理 / 租赁账单 / 查看。")
    add_heading_cn(doc, "4.1 账单信息", 2)
    add_para_cn(
        doc,
        "展示合同编码、合同类型、项目名称、客户名称、交车任务编码、账单编码、账单期数、"
        "账单开始日期、账单结束日期（与列表规则一致）。",
    )
    add_heading_cn(doc, "4.2 账单明细汇总", 2)
    add_bullet(
        doc,
        "应收款总额（可点击）：展开包含「应收月租金合计、应收保证金合计、应收服务费合计」，"
        "首期账单另含「氢费预付款应收」。",
    )
    add_bullet(
        doc,
        "实收款总额（可点击）：包含实收租金、保证金、实收服务费、租金减免、服务费减免等；"
        "首期另含氢费预收与氢费减免。",
    )
    add_bullet(doc, "开票总额：与实收口径一致但不含「应收车辆保证金」，用于开票参考。")
    add_heading_cn(doc, "4.3 车辆明细表", 2)
    add_para_cn(
        doc,
        "每车一行：序号、品牌、型号、车牌、应收/实收月租金、租金备注、减免金额及备注、减免证明（可点预览）、"
        "应收保证金（首期按合同，非首期为 0）、服务费项目（点「管理」弹出各服务项目应收/实收/减免/备注）、"
        "应收服务费与实收服务费合计。",
    )
    add_heading_cn(doc, "4.4 氢费预付款（仅首期）", 2)
    add_para_cn(doc, "首期账单在表格下方展示氢费预付款应收、实收、减免金额及备注；第二期起整块不显示。")
    add_heading_cn(doc, "4.5 返回", 2)
    add_para_cn(doc, "点击「返回」回到租赁账单列表。")

    add_heading_cn(doc, "五、收费明细页（编辑提交）", 1)
    add_para_cn(doc, "面包屑：业务管理 / 租赁账单 / 收费明细。由列表子表「收费明细」进入。")
    add_heading_cn(doc, "5.1 账单信息与汇总", 2)
    add_para_cn(doc, "顶部「账单信息」与查看页相同。汇总区同样可点击「应收款总额」「实收款总额」查看分项；「开票金额」规则同查看页。")
    add_heading_cn(doc, "5.2 可编辑内容", 2)
    add_bullet(doc, "实收车辆月租金：输入框，两位小数，默认与应收一致，可按实收修改。")
    add_bullet(doc, "车辆租金备注、减免金额、减免备注：文本或金额输入。")
    add_bullet(doc, "减免证明：附件上传（多文件），单行展示文件名，可预览、可删除。")
    add_bullet(doc, "服务费项目：点「管理」，在气泡内维护每项「实收费用」（必填 *）、「减免费用」、「备注」；实收服务费自动按实收费用汇总。")
    add_bullet(doc, "氢费预付款（仅首期）：实收金额、减免金额、减免备注可编辑；应收金额为合同带出只读。")
    add_heading_cn(doc, "5.3 底部按钮", 2)
    add_bullet(doc, "提交：弹出「请确认账单金额无误」，确认后提交成功并返回（原型逻辑）。")
    add_bullet(doc, "保存：保存当前填写，不做完整校验，保存后返回。")
    add_bullet(doc, "取消：提示未保存将丢失修改，确认后返回列表。")

    add_heading_cn(doc, "六、操作建议流程（培训口径）", 1)
    add_para_cn(doc, "1）在列表用筛选定位合同 → 展开子表找到目标期数 → 看「状态」。")
    add_para_cn(doc, "2）先点「查看」核对应收与合同是否一致。")
    add_para_cn(doc, "3）若为期初「待提交」，点「收费明细」补全实收、减免、服务费与附件 → 保存或提交。")
    add_para_cn(doc, "4）提交后仅能通过「查看」复核；列表中可继续维护氢费成本、其他成本（与收费明细提交状态独立，以实际上线规则为准）。")

    add_heading_cn(doc, "七、文档信息", 1)
    add_para_cn(doc, "版本：与原型标注一致（租赁账单 2026-03-10；查看/收费明细 2026-03-11）。生成工具：项目内 web端/业务管理/文档/_build_租赁账单手册.py")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print("Saved:", out_path)


if __name__ == "__main__":
    desktop = Path.home() / "Desktop"
    build(desktop / "租赁账单功能说明.docx")
