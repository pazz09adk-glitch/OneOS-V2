# -*- coding: utf-8 -*-
"""生成「还车应结款」用户操作说明 Word 文档及原型配图（脚本可重复运行）"""
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:
    Image = None


ROOT = Path(__file__).resolve().parent
OUT_DOC = ROOT / "还车应结款-用户操作说明.docx"
LIST_SCREENSHOT = Path(
    "/Users/sylvawong/.cursor/projects/Users-sylvawong-Desktop-CURSOR-ONE-OS/assets/"
    "2230B49A-9740-471C-ADE4-FC7723DC4CA9_4_5005_c-70f13b01-98b2-42e3-a4f5-6824f02f1838.png"
)


def set_cell_shading(cell, fill_hex):
    """Word 表格单元格背景色"""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_cn(doc, text, level):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "PingFang SC"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    return h


def add_para_cn(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "PingFang SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    run.bold = bold
    return p


def make_prototype_figure(path: Path, title: str, lines: list):
    """根据原型结构生成示意配图（非系统实拍，便于培训对照 JSX页面区块）"""
    if Image is None:
        return False
    w, h = 900, 520
    img = Image.new("RGB", (w, h), (250, 250, 250))
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, w - 1, h - 1], outline=(200, 200, 200), width=2)
    try:
        font_title = ImageFont.truetype("/System/Library/Fonts/PingFang.ttc", 22, index=1)
        font_body = ImageFont.truetype("/System/Library/Fonts/PingFang.ttc", 17, index=1)
        font_small = ImageFont.truetype("/System/Library/Fonts/PingFang.ttc", 14, index=1)
    except OSError:
        font_title = font_body = font_small = ImageFont.load_default()

    y = 20
    draw.text((24, y), title, fill=(17, 24, 39), font=font_title)
    y += 44
    draw.line([(24, y), (w - 24, y)], fill=(217, 217, 217), width=1)
    y += 16
    for line in lines:
        draw.text((32, y), line, fill=(51, 51, 51), font=font_body)
        y += 28
    note = "说明：本图为原型页面结构示意，正式文档可替换为系统实拍截图。"
    draw.text((24, h - 36), note, fill=(120, 120, 120), font=font_small)
    img.save(path, "PNG")
    return True


def build():
    ROOT.mkdir(parents=True, exist_ok=True)
    fig_view = ROOT / "配图-还车应结款-查看页-原型结构.png"
    fig_fee = ROOT / "配图-还车应结款-费用明细-原型结构.png"

    make_prototype_figure(
        fig_view,
        "还车应结款 · 查看页（原型 还车应结款-查看.jsx）",
        [
            "面包屑：财务管理 / 还车应结款 / 查看",
            "卡片① 还车车辆明细表（车牌、合同、项目、客户、交还车时间）",
            "  · 易损保 / 轮胎保 / 养护保（是/否 + 提示图标说明）",
            "卡片② 还车费用明细",
            "  · 统计：保证金、待结算、应退还、应补缴（可点开分项）",
            "  · 业务服务组 / 能源组 / 运维部 / 安全组（只读展示）",
            "底部：返回列表",
        ],
    )
    make_prototype_figure(
        fig_fee,
        "还车应结款 · 费用明细页（原型 还车应结款-费用明细.jsx）",
        [
            "面包屑：财务管理 / 还车应结款（含「查看需求说明」）",
            "卡片 还车车辆明细（同上）",
            "卡片 还车费用明细：顶部四统计 + 可折叠四组",
            "  · 业务服务组：固定费用行 + 可增删行；车辆租金三块",
            "  · 能源采购组：氢量差/交还车氢量/单价/氢电费/预付款退费",
            "  · 运维部：清洗保养维修等；轮胎磨损说明气泡；无忧包减免",
            "  · 安全组：违章清单 + 事故清单；保存/提交/撤回",
            "底栏：提交审核（15天倒计时+四组已提交） / 取消",
        ],
    )

    doc = Document()
    sect = doc.sections[0]
    sect.page_height = Cm(29.7)
    sect.page_width = Cm(21.0)
    sect.left_margin = Cm(2.2)
    sect.right_margin = Cm(2.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("数字化资产 ONEOS 运管平台\n还车应结款 · 用户操作说明（培训版）")
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "PingFang SC"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    add_para_cn(
        doc,
        "适用范围：财务管理模块下的「还车应结款」列表、只读查看页、费用明细协作页。"
        "配图说明：图1 为同类财务列表界面实拍（布局与还车应结款列表一致）；"
        "图 2、图 3 为依据原型 JSX 页面结构整理的示意配图，可在定稿前替换为系统截图。",
    )

    add_heading_cn(doc, "一、模块作用与入口", 1)
    add_para_cn(
        doc,
        "用于车辆还车后汇总安全、业务服务、运维、能源等数据，形成待结算金额并进入审批/账单流程。"
        "入口：左侧菜单 财务管理 → 还车应结款。",
    )

    add_heading_cn(doc, "二、列表页：筛选与查询", 1)
    add_para_cn(doc, "筛选区支持：合同编号、客户名称、项目名称（可展开后：车牌号、还车时间范围、审批状态）。按钮：重置、查询。")
    tbl = doc.add_table(rows=8, cols=2)
    tbl.style = "Table Grid"
    hdr = ["条件", "说明"]
    for j, t in enumerate(hdr):
        tbl.rows[0].cells[j].text = t
        set_cell_shading(tbl.rows[0].cells[j], "E6F4EA")
    rows_data = [
        ("合同编号", "下拉 + 输入模糊匹配"),
        ("客户名称", "同上"),
        ("项目名称", "同上"),
        ("车牌号", "展开后，同上"),
        ("还车时间", "起止日期，精确到日"),
        ("审批状态", "待提交/待审批/审批中/审批完成/审批驳回/撤回"),
        ("导出", "列表卡片右上角导出当前结果"),
    ]
    for i, (a, b) in enumerate(rows_data, start=1):
        tbl.rows[i].cells[0].text = a
        tbl.rows[i].cells[1].text = b

    add_heading_cn(doc, "三、列表页：表格字段与审批状态", 1)
    add_para_cn(doc, "提交情况四列：安全组、业务服务组、运维组、能源组。绿点=已提交，灰点=未提交，后接提交人姓名。")
    add_para_cn(doc, "审批状态含义：待提交=四组未齐且未提审批；待审批=已提审无人处理；审批中=有节点已通过未完；审批完成=流程结束；审批驳回=任一节点的驳回；撤回=终审前主动撤回。")

    add_heading_cn(doc, "四、列表页：操作按钮", 1)
    ops = [
        "查看：进入只读「查看」页。",
        "生成账单：仅「待审批」显示；弹窗账单并可打印（需允许浏览器弹窗）。",
        "费用明细：待审批/审批中/审批完成 时隐藏；其余可编辑阶段进入费用明细。",
        "撤回：仅待审批、审批中显示；确认后状态改为撤回。",
    ]
    for o in ops:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(o)
        r.font.name = "PingFang SC"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        r.font.size = Pt(11)

    add_heading_cn(doc, "五、配图", 1)

    p_cap1 = doc.add_paragraph()
    p_cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_cap1.add_run(
        "图 1  财务管理 · 列表页布局参考（实拍；与「还车应结款」列表同一套：侧栏、筛选、表格、分页）"
    )
    r1.font.size = Pt(10)
    r1.font.name = "PingFang SC"
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    r1.italic = True

    if LIST_SCREENSHOT.is_file():
        doc.add_picture(str(LIST_SCREENSHOT), width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_para_cn(doc, "（未找到图1源文件，请将财务列表截图置于文档此处）", bold=True)

    p_cap2 = doc.add_paragraph()
    p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_cap2.add_run("图 2  还车应结款 · 查看页（原型结构示意）")
    r2.font.size = Pt(10)
    r2.font.name = "PingFang SC"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    r2.italic = True
    if fig_view.is_file():
        doc.add_picture(str(fig_view), width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_cap3 = doc.add_paragraph()
    p_cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p_cap3.add_run("图 3  还车应结款 · 费用明细页（原型结构示意）")
    r3.font.size = Pt(10)
    r3.font.name = "PingFang SC"
    r3._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    r3.italic = True
    if fig_fee.is_file():
        doc.add_picture(str(fig_fee), width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading_cn(doc, "六、查看页（只读）要点", 1)
    add_para_cn(
        doc,
        "车辆明细含三项保险及提示文案。费用区四张统计卡可点击展开分项。"
        "业务服务组、能源组、运维部、安全组为只读表格；轮胎磨损可悬停查看逐胎明细。底部返回列表。",
    )

    add_heading_cn(doc, "七、费用明细页：各组分工", 1)
    add_para_cn(
        doc,
        "业务服务组：5 项固定费用 + 可新增行；车辆租金（已收/实际/应退，实际租金默认按日折算可手改）。保存/提交/撤回；已提交后锁定至撤回。",
    )
    add_para_cn(
        doc,
        "能源采购组：交还车氢量、单价只读；交车氢量大于还车时氢量差补缴自动算；氢费/电费/预付款退费可填；最后一辆车红色提示。",
    )
    add_para_cn(
        doc,
        "运维部：清洗、保养、维修、车损等；证件丢失按规则自动计价；送/接车服务费禁用反写；轮胎磨损与胎纹合计规则见需求说明；无忧包减免与三项保险挂钩。",
    )
    add_para_cn(
        doc,
        "安全组：违章清单、事故清单；确认后提交。",
    )

    add_heading_cn(doc, "八、金额计算（口径）", 1)
    add_para_cn(
        doc,
        "待结算总额 = 业务服务组费用合计 + 车辆应退租金 + 氢量差补缴 + 氢费补缴 + 电费补缴 − 预付款退费 + 运维部费用合计。",
    )
    add_para_cn(
        doc,
        "应退还总额：保证金 − 待结算 为正时取该值。应补缴总额：保证金 − 待结算 为负时取绝对值。",
    )

    add_heading_cn(doc, "九、提交审核条件", 1)
    add_para_cn(
        doc,
        "原型规则：单据生成后满 15 天倒计时结束，且业务服务组、能源采购组、运维部、安全组均为「已提交」，"
        "底部「提交审核」才可点；否则按钮禁用并显示剩余天/小时。提交前系统校验各必填金额与无忧包减免等。",
    )

    add_heading_cn(doc, "十、修订记录", 1)
    add_para_cn(doc, "文档版本：V1.0 生成依据：web端/财务管理/还车应结款.jsx、还车应结款-查看.jsx、还车应结款-费用明细.jsx 内嵌需求说明。")

    doc.save(OUT_DOC)
    print("Wrote:", OUT_DOC)


if __name__ == "__main__":
    build()
